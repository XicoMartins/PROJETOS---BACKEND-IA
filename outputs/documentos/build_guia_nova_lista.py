from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PANEL_ROOT = Path(
    r"S:\PROJETOS EM ANDAMENTO\PAINEL DE CONTROLE MTECH\PROGRAMAS"
    r"\PROJETOS - PAINEL PRODUÇÃO IA"
)
LOGO = PANEL_ROOT / "brand_logo.png"
OUTPUT = Path(__file__).resolve().parent / "GUIA_NOVA_LISTA_PROCESSO_MTECH.docx"

TEAL = "00A3A3"
TEAL_LIGHT = "7CD4D4"
PETROL = "0B3A44"
INK = "172027"
MUTED = "60747A"
PALE = "E8F7F7"
PALE_BLUE = "DDF0F2"
LIGHT_GRAY = "F3F6F7"
WHITE = "FFFFFF"
RED = "A33A3A"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, color="C5D9DC", size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def mark_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(
    run,
    *,
    name="Calibri",
    size=11,
    color=INK,
    bold=None,
    italic=None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, side="bottom", color=TEAL, size=12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 27, PETROL, 0, 8),
        ("Subtitle", 12.5, MUTED, 0, 18),
        ("Heading 1", 16, PETROL, 18, 10),
        ("Heading 2", 13, TEAL, 14, 7),
        ("Heading 3", 11.5, PETROL, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    def add_definition(abstract_id: int, num_id: int, fmt: str, text: str, color: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), color)
        r_pr.append(clr)
        bold = OxmlElement("w:b")
        r_pr.append(bold)
        level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    existing_abs = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    base_abs = max(existing_abs, default=0) + 10
    base_num = max(existing_num, default=0) + 10
    add_definition(base_abs, base_num, "decimal", "%1.", TEAL)
    add_definition(base_abs + 1, base_num + 1, "bullet", "•", TEAL)
    return base_num, base_num + 1


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_numbered_item(doc, num_id: int, lead: str, detail: str) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(lead)
    set_run_font(run, bold=True, color=PETROL)
    run = p.add_run(detail)
    set_run_font(run)


def add_bullet(doc, num_id: int, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(3)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True, color=PETROL)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        set_run_font(p.add_run(text))


def add_callout(doc, label: str, text: str, *, fill=PALE, color=PETROL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, side="left", color=TEAL, size=18)
    run = p.add_run(f"  {label.upper()}  ")
    set_run_font(run, size=10.5, bold=True, color=color)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [2700, 6660], indent_dxa=0)
    logo = table.rows[0].cells[0].paragraphs[0].add_run().add_picture(
        str(LOGO), width=Inches(1.45)
    )
    logo._inline.docPr.set("descr", "Logotipo MTECH Displays")
    logo._inline.docPr.set("title", "MTECH Displays")
    mark_table_header(table.rows[0])
    right = table.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(
        right.add_run("GUIA OPERACIONAL  |  AUTOMAÇÃO DE PROCESSOS"),
        size=8.5,
        color=MUTED,
        bold=True,
    )
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=40, end=0)
    set_paragraph_border(header.paragraphs[-1], color=TEAL, size=10)

    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph_border(p, side="top", color=TEAL_LIGHT, size=6)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(p.add_run("MTECH Displays  •  Uso interno  |  "), size=8.5, color=MUTED)
    set_run_font(p.add_run("Página "), size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_quick_comparison(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680])
    headers = (("FAÇA", TEAL), ("NÃO FAÇA", PETROL))
    for idx, (label, fill) in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(label), size=10.5, color=WHITE, bold=True)
    mark_table_header(table.rows[0])

    rows = [
        (
            "Salvar uma cópia em .xlsx e fechar o Excel.",
            "Colocar arquivo aberto ou temporário (~$...).",
        ),
        (
            "Usar um nome novo e descritivo para a lista.",
            "Reutilizar nome que já existe na base.",
        ),
        (
            "Colocar somente na pasta de entrada.",
            "Copiar manualmente para planilhas ou para o painel.",
        ),
        (
            "Aguardar a automação gerar os IDs.",
            "Preencher ou reaproveitar PROCESSO_ID por conta própria.",
        ),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((left, right)):
            set_cell_border(cells[idx])
            set_cell_shading(cells[idx], WHITE if idx == 0 else LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text), size=10)
    set_table_geometry(table, [4680, 4680])


def add_path_block(doc: Document, path: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, PETROL)
    run = p.add_run(path)
    set_run_font(run, name="Consolas", size=8.5, color=WHITE, bold=True)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)

    configure_styles(doc)
    num_id, bullet_id = create_numbering(doc)
    add_header_footer(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(2)
    set_run_font(
        kicker.add_run("ROTINA DO ANALISTA"),
        size=9.5,
        color=TEAL,
        bold=True,
    )

    title = doc.add_paragraph(style="Title")
    set_paragraph_border(title, color=TEAL, size=14)
    title.add_run("Nova lista de processo no backend")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Passo a passo para cadastrar uma nova planilha com geração automática "
        "de PROCESSO_ID, QR Codes e publicação nos painéis."
    )

    add_callout(
        doc,
        "Regra principal",
        "O analista coloca a planilha somente na pasta de entrada. "
        "Não é necessário copiar para a base, para o painel ou para o GitHub.",
    )

    doc.add_paragraph("1. Antes de enviar", style="Heading 1")
    add_bullet(
        doc,
        bullet_id,
        "Use o modelo Excel atual da lista de processo e salve no formato .xlsx.",
    )
    add_bullet(
        doc,
        bullet_id,
        "Confirme as colunas obrigatórias: CLIENTE, ACABADO, FERRAMENTAL e PROCESSO.",
        bold_lead="Confirme as colunas obrigatórias:",
    )
    add_bullet(
        doc,
        bullet_id,
        "PROCESSO_ID pode não existir ou estar vazio; a automação preencherá a sequência.",
        bold_lead="PROCESSO_ID",
    )
    add_bullet(
        doc,
        bullet_id,
        "Use um nome de arquivo novo, claro e único. Exemplo: LISTA DE PROCESSO RACK MODELO X.xlsx.",
    )
    add_bullet(
        doc,
        bullet_id,
        "Feche completamente o arquivo no Excel antes de copiá-lo.",
        bold_lead="Feche completamente",
    )

    doc.add_paragraph("2. Enviar a planilha", style="Heading 1")
    add_numbered_item(
        doc,
        num_id,
        "Abra a pasta de entrada. ",
        "Para lista de produção, use exatamente o caminho abaixo.",
    )
    add_path_block(
        doc,
        r"S:\PROJETOS EM ANDAMENTO\PAINEL DE CONTROLE MTECH\PROGRAMAS"
        r"\PROJETOS---BACKEND-IA\automacao_qr\entrada\producao",
    )
    add_numbered_item(
        doc,
        num_id,
        "Copie uma única vez. ",
        "Cole o arquivo .xlsx fechado nessa pasta.",
    )
    add_numbered_item(
        doc,
        num_id,
        "Aguarde a automação. ",
        "A verificação acontece automaticamente, normalmente em até 1 minuto.",
    )
    add_numbered_item(
        doc,
        num_id,
        "Não mova nem renomeie durante o processamento. ",
        "O arquivo sairá sozinho da entrada quando terminar.",
    )

    doc.add_page_break()

    doc.add_paragraph("3. O que acontece automaticamente", style="Heading 1")
    flow = [
        ("VALIDAÇÃO", "estrutura, campos obrigatórios, nome e duplicidades"),
        ("IDENTIFICAÇÃO", "criação da coluna PROCESSO_ID e sequência global"),
        ("QR CODES", "um arquivo PNG para cada processo"),
        ("PUBLICAÇÃO", "backend e painel de produção"),
        ("GITHUB", "commit e push nos dois repositórios"),
        ("ARQUIVAMENTO", "original movido para a pasta processados"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2100, 7260])
    for idx, label in enumerate(("ETAPA", "AÇÃO AUTOMÁTICA")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PETROL)
        set_cell_border(cell, color=PETROL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(label), size=9.5, color=WHITE, bold=True)
    mark_table_header(table.rows[0])
    for idx, (label, detail) in enumerate(flow):
        cells = table.add_row().cells
        set_cell_shading(cells[0], TEAL if idx % 2 == 0 else PETROL)
        set_cell_shading(cells[1], PALE if idx % 2 == 0 else LIGHT_GRAY)
        for cell in cells:
            set_cell_border(cell, color="B8D1D4")
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(label), size=9.5, color=WHITE, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(detail), size=10.2)
    set_table_geometry(table, [2100, 7260])

    doc.add_paragraph("4. Como confirmar que deu certo", style="Heading 1")
    checks = [
        "O arquivo não está mais na pasta de entrada.",
        "Existe uma cópia com PROCESSO_ID preenchido na pasta backend\\planilhas.",
        "A mesma planilha aparece na pasta planilhas do painel de produção.",
        "Foi criada uma pasta com os QR Codes em qrcodes_processos\\base_completa.",
        "O original está arquivado em automacao_qr\\processados, organizado por data.",
    ]
    for item in checks:
        add_bullet(doc, bullet_id, item)

    add_callout(
        doc,
        "Tempo esperado",
        "A rotina verifica a entrada a cada minuto. Planilhas maiores e o envio ao "
        "GitHub podem acrescentar alguns segundos.",
        fill=PALE_BLUE,
    )

    doc.add_paragraph("5. Se houver erro", style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("A planilha será movida para "), size=10.5)
    set_run_font(
        p.add_run(r"automacao_qr\rejeitados"),
        name="Consolas",
        size=9.5,
        color=RED,
        bold=True,
    )
    set_run_font(
        p.add_run(" e receberá um arquivo .erro.txt com a explicação."),
        size=10.5,
    )
    add_bullet(
        doc,
        bullet_id,
        "Leia o .erro.txt, corrija a planilha e coloque novamente na entrada.",
    )
    add_bullet(
        doc,
        bullet_id,
        "Se a mensagem informar nome já existente, não sobrescreva a base; procure o responsável pelo sistema.",
    )
    add_bullet(
        doc,
        bullet_id,
        "Se o repositório estiver com alterações pendentes, avise o responsável técnico.",
    )

    doc.add_paragraph("Resumo rápido", style="Heading 1")
    add_quick_comparison(doc)

    add_callout(
        doc,
        "Em caso de dúvida",
        "Não copie a planilha manualmente para outras pastas. Preserve o arquivo "
        "rejeitado e encaminhe a mensagem do .erro.txt ao responsável técnico.",
        fill=LIGHT_GRAY,
    )

    core = doc.core_properties
    core.title = "Guia para subir nova lista de processo no backend"
    core.subject = "Procedimento operacional MTECH"
    core.author = "MTECH Displays"
    core.keywords = "MTECH, lista de processo, backend, QR Code, PROCESSO_ID"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
